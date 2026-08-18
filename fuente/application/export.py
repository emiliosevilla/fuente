"""Canonical note export from NoteDocument (Task 6.4)."""
from __future__ import annotations

import base64
import html
import io
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Literal
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document

from fuente.application.notes import NotesApplicationService
from fuente.domain.documents import NoteDocument
from fuente.domain.frontmatter import parse_frontmatter, serialize_frontmatter
from fuente.domain.paths import AuthorizedPathResolver
from fuente.infrastructure.atomic_files import atomic_write_text

ExportFormat = Literal["markdown", "docx", "pdf"]

_FORMAT_EXTENSIONS: dict[ExportFormat, str] = {
    "markdown": ".md",
    "docx": ".docx",
    "pdf": ".pdf",
}

_FORMAT_MIME: dict[ExportFormat, str] = {
    "markdown": "text/markdown;charset=utf-8",
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
    "pdf": "application/pdf",
}

_DOCX_FENCE_RE = re.compile(r"^\s*(`{3,}|~{3,})(.*)$")
_DOCX_ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)


class ExportProjectionError(Exception):
    """Known failure while projecting a canonical note into an export payload."""

    code = "export_projection_failed"


class ExportFileExistsError(ExportProjectionError, FileExistsError):
    """Raised when an export target exists and overwrite was not confirmed."""

    code = "export_file_exists"

    def __init__(self, destination: str) -> None:
        self.destination = destination
        super().__init__(f"Export destination already exists: {destination}")


class UnsupportedExportFormatError(ExportProjectionError, ValueError):
    """Raised when the UI requests an unsupported export format."""

    code = "unsupported_export_format"

    def __init__(self, export_format: str) -> None:
        self.export_format = export_format
        super().__init__(f"Unsupported export format: {export_format}")


@dataclass(frozen=True)
class ExportPayload:
    """Prepared export artifact derived from the canonical note document."""

    format: ExportFormat
    filename: str
    source: str
    content: str | None = None
    content_bytes: bytes | None = None
    content_type: str = ""
    mode: str = "download"
    label: str = ""
    print_html: str = ""

    def as_dict(self) -> dict:
        payload: dict = {
            "format": self.format,
            "filename": self.filename,
            "source": self.source,
            "content_type": self.content_type or _FORMAT_MIME[self.format],
            "mode": self.mode,
        }
        if self.label:
            payload["label"] = self.label
        if self.content is not None:
            payload["content"] = self.content
        if self.content_bytes is not None:
            payload["content_base64"] = base64.b64encode(self.content_bytes).decode(
                "ascii"
            )
        if self.print_html:
            payload["print_html"] = self.print_html
        return payload


class ExportApplicationService:
    """Export notes from canonical NoteDocument, never from rendered DOM."""

    def __init__(
        self,
        *,
        notes_service: NotesApplicationService,
        path_resolver: AuthorizedPathResolver,
    ) -> None:
        self.notes_service = notes_service
        self.path_resolver = path_resolver

    def prepare_download(self, document_id: str, export_format: str) -> ExportPayload:
        """Build a deterministic export payload for client-side download."""
        note = self.notes_service.get_note(document_id)
        self.notes_service.require_published_output(note)
        normalized = self._normalize_format(export_format)
        if normalized == "markdown":
            return self._markdown_payload(note)
        if normalized == "docx":
            return self._docx_payload(note)
        return self._pdf_payload(note)

    def write_export(
        self,
        document_id: str,
        export_format: str,
        destination_relative_path: str,
        *,
        confirm_overwrite: bool = False,
    ) -> dict:
        """Write an export to an authorized vault-relative destination."""
        normalized = self._normalize_format(export_format)
        destination = self._resolve_destination(
            destination_relative_path, export_format=normalized
        )
        if destination.exists() and not confirm_overwrite:
            raise ExportFileExistsError(destination_relative_path)

        payload = self.prepare_download(document_id, normalized)
        if normalized == "markdown":
            atomic_write_text(destination, payload.content or "")
        elif normalized == "docx":
            destination.parent.mkdir(parents=True, exist_ok=True)
            destination.write_bytes(payload.content_bytes or b"")
        else:
            raise UnsupportedExportFormatError(normalized)

        return {
            "status": "exported",
            "format": normalized,
            "path": self._vault_relative(destination),
            "filename": destination.name,
        }

    @staticmethod
    def _normalize_format(export_format: str) -> ExportFormat:
        cleaned = (export_format or "").strip().lower()
        aliases = {"word": "docx", "md": "markdown"}
        cleaned = aliases.get(cleaned, cleaned)
        if cleaned not in _FORMAT_EXTENSIONS:
            raise UnsupportedExportFormatError(export_format)
        return cleaned  # type: ignore[return-value]

    @staticmethod
    def resolve_destination(
        path_resolver: AuthorizedPathResolver,
        destination_relative_path: str,
        *,
        export_format: ExportFormat,
    ) -> Path:
        allowed = {_FORMAT_EXTENSIONS[export_format]}
        return path_resolver.resolve(
            destination_relative_path,
            root_name="output",
            allowed_extensions=allowed,
        )

    def _resolve_destination(
        self, destination_relative_path: str, *, export_format: ExportFormat
    ) -> Path:
        return self.resolve_destination(
            self.path_resolver,
            destination_relative_path,
            export_format=export_format,
        )

    def _vault_relative(self, path: Path) -> str:
        vault_root = self.path_resolver.roots["vault"]
        return path.resolve().relative_to(vault_root.resolve()).as_posix()

    def _safe_filename(self, note: NoteDocument, suffix: str) -> str:
        stem = Path(note.relative_path).stem or "nota_fuente"
        cleaned = re.sub(r"[^\w.\- ]+", "_", stem).strip("._ ") or "nota_fuente"
        return f"{cleaned}{suffix}"

    def _markdown_payload(self, note: NoteDocument) -> ExportPayload:
        return ExportPayload(
            format="markdown",
            filename=self._safe_filename(note, ".md"),
            source="canonical",
            content=note.to_markdown(),
            content_type=_FORMAT_MIME["markdown"],
            mode="download",
        )

    def _docx_payload(self, note: NoteDocument) -> ExportPayload:
        return ExportPayload(
            format="docx",
            filename=self._safe_filename(note, ".docx"),
            source="canonical",
            content_bytes=self._render_docx(note),
            content_type=_FORMAT_MIME["docx"],
            mode="download",
        )

    def _pdf_payload(self, note: NoteDocument) -> ExportPayload:
        return ExportPayload(
            format="pdf",
            filename=self._safe_filename(note, ".pdf"),
            source="canonical",
            content_type=_FORMAT_MIME["pdf"],
            mode="user_assisted_print",
            label="PDF (impresión asistida, contenido canónico completo)",
            print_html=self._render_print_html(note),
        )

    @staticmethod
    def _render_docx(note: NoteDocument) -> bytes:
        canonical = note.to_markdown()
        metadata, body = parse_frontmatter(canonical)
        document = Document()
        document.add_heading(note.title or Path(note.relative_path).stem, level=0)
        document.add_paragraph(f"Ruta: {note.relative_path}")
        document.add_paragraph("---")

        metadata_table = document.add_table(rows=0, cols=2)
        metadata_table.style = "Table Grid"
        for key in sorted(metadata):
            cells = metadata_table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = ExportApplicationService._docx_metadata_value(
                metadata[key]
            )

        paragraph_lines: list[str] = []
        code_lines: list[str] = []
        code_fence: str | None = None

        def flush_paragraph() -> None:
            if paragraph_lines:
                document.add_paragraph("\n".join(paragraph_lines))
                paragraph_lines.clear()

        def flush_code() -> None:
            code_paragraph = document.add_paragraph(style="No Spacing")
            code_run = code_paragraph.add_run("\n".join(code_lines))
            code_run.font.name = "Courier New"
            code_lines.clear()

        def is_closing_fence(line: str, opening_fence: str) -> bool:
            stripped = line.strip()
            return (
                bool(stripped)
                and stripped[0] == opening_fence[0]
                and len(stripped) >= len(opening_fence)
                and set(stripped) == {opening_fence[0]}
            )

        for line in body.splitlines():
            if code_fence is not None:
                if is_closing_fence(line, code_fence):
                    flush_code()
                    code_fence = None
                else:
                    code_lines.append(line)
                continue

            fence = _DOCX_FENCE_RE.match(line)
            if fence:
                flush_paragraph()
                code_fence = fence.group(1)
                continue

            if not line.strip():
                flush_paragraph()
                continue

            heading = re.match(r"^(#{1,3})\s+(.*)$", line)
            if heading:
                flush_paragraph()
                document.add_heading(heading.group(2), level=len(heading.group(1)))
                continue

            if line.startswith("- "):
                flush_paragraph()
                document.add_paragraph(line[2:], style="List Bullet")
                continue

            numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
            if numbered:
                flush_paragraph()
                document.add_paragraph(numbered.group(2), style="List Number")
                continue

            paragraph_lines.append(line)

        flush_paragraph()
        if code_fence is not None:
            flush_code()
        buffer = io.BytesIO()
        document.save(buffer)
        return ExportApplicationService._canonicalize_docx(buffer.getvalue())

    @staticmethod
    def _canonicalize_docx(raw: bytes) -> bytes:
        source = io.BytesIO(raw)
        target = io.BytesIO()
        with ZipFile(source, "r") as archive, ZipFile(
            target, "w", compression=ZIP_DEFLATED, compresslevel=9
        ) as canonical:
            for name in sorted(archive.namelist()):
                info = ZipInfo(name, _DOCX_ZIP_TIMESTAMP)
                info.create_system = 3
                info.compress_type = ZIP_DEFLATED
                info.compress_level = 9
                info.external_attr = 0o600 << 16
                canonical.writestr(info, archive.read(name))
        return target.getvalue()

    @staticmethod
    def _docx_metadata_value(value: object) -> str:
        if isinstance(value, (list, dict)):
            return json.dumps(value, ensure_ascii=False, sort_keys=True)
        return str(value)

    def _render_print_html(self, note: NoteDocument) -> str:
        """Render user-assisted print HTML from canonical frontmatter + body."""
        title = html.escape(note.title or Path(note.relative_path).stem)
        relative_path = html.escape(note.relative_path)
        canonical = note.to_markdown()
        metadata, body = parse_frontmatter(canonical)
        frontmatter_block = serialize_frontmatter(metadata).rstrip()
        frontmatter_html = (
            f"<pre class='pdf-frontmatter'>{html.escape(frontmatter_block)}</pre>"
        )
        body_html = self._markdown_body_to_print_html(body)
        return (
            "<!DOCTYPE html><html><head><meta charset='utf-8'>"
            f"<title>{title} — Exportación PDF</title>"
            "<style>"
            "@page { size: A4; margin: 20mm; }"
            "body { font-family: Arial, Helvetica, sans-serif; font-size: 12pt; "
            "line-height: 1.55; color: #111111; margin: 0; padding: 0; }"
            ".pdf-header { border-bottom: 1.5px solid #888888; padding-bottom: 8px; "
            "margin-bottom: 22px; }"
            ".pdf-header-title { font-size: 9pt; font-weight: bold; color: #222222; }"
            ".pdf-header-path { font-size: 7.8pt; color: #555555; word-break: break-all; }"
            ".pdf-frontmatter { white-space: pre-wrap; font-family: 'Courier New', monospace; "
            "font-size: 9pt; background: #f7f7f7; border: 1px solid #cccccc; "
            "padding: 10px; margin-bottom: 18px; }"
            "h1 { font-size: 20pt; margin-top: 0; }"
            "h2 { font-size: 16pt; }"
            "h3 { font-size: 14pt; }"
            "pre { white-space: pre-wrap; font-family: 'Courier New', monospace; }"
            "</style></head><body>"
            "<div class='pdf-header'>"
            f"<div class='pdf-header-title'>Nota: {title}</div>"
            f"<div class='pdf-header-path'>Ruta: {relative_path}</div>"
            "</div>"
            f"{frontmatter_html}"
            f"{body_html}"
            "<p><em>Exportación asistida: guarda como PDF desde el diálogo de impresión. "
            "Incluye frontmatter y cuerpo canónicos.</em></p>"
            "</body></html>"
        )

    @staticmethod
    def _markdown_body_to_print_html(body_markdown: str) -> str:
        blocks: list[str] = []
        for line in body_markdown.splitlines():
            escaped = html.escape(line)
            if line.startswith("# "):
                blocks.append(f"<h1>{html.escape(line[2:])}</h1>")
            elif line.startswith("## "):
                blocks.append(f"<h2>{html.escape(line[3:])}</h2>")
            elif line.startswith("### "):
                blocks.append(f"<h3>{html.escape(line[4:])}</h3>")
            elif line.strip() == "":
                continue
            else:
                blocks.append(f"<p>{escaped}</p>")
        if not blocks:
            blocks.append("<p></p>")
        return "".join(blocks)
