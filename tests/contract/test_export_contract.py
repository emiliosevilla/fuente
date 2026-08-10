"""Export ↔ canonical document contract matrix (Task 8.3)."""
from __future__ import annotations

import base64
import io
import inspect
import json

from docx import Document
import pytest

from funes.application.export import ExportApplicationService
from funes.application.notes import NotesApplicationService
from funes.control_console import FunesConsoleBackend
from funes.domain.errors import NoteRevisionConflictError
from funes.domain.frontmatter import parse_frontmatter
from funes.domain.metadata_form import MetadataValidationError
from funes.domain.paths import AuthorizedPathResolver
from funes.infrastructure.sqlite_store import JobStore
from funes.ui.bridge import FunesPyWebViewApi

from tests.contract.conftest import CONSOLA_HTML, write_note_under_theme

THEME = "Derecho_Civil"
ISSUE = "Contratos"


@pytest.fixture
def export_stack(temp_vault_manager):
    resolver = AuthorizedPathResolver(
        vault_root=temp_vault_manager.config.vault_path,
        output=temp_vault_manager.output_dir,
        input=temp_vault_manager.input_dir,
        dirty=temp_vault_manager.dirty_dir,
        clean=temp_vault_manager.clean_dir,
        quarantine=temp_vault_manager.quarantine_dir,
    )
    store = JobStore(temp_vault_manager.config.vault_path)
    notes = NotesApplicationService(
        vault=temp_vault_manager,
        path_resolver=resolver,
        job_store=store,
        chroma_store=None,
    )
    export_service = ExportApplicationService(
        notes_service=notes,
        path_resolver=resolver,
    )
    try:
        yield export_service, notes, temp_vault_manager
    finally:
        store.close()


def _prepare_approved_note(vault_manager) -> str:
    vault_manager.create_theme(THEME)
    vault_manager.create_issue_in_theme(ISSUE)
    document_id, _ = write_note_under_theme(
        vault_manager,
        theme=THEME,
        issue=ISSUE,
        title="Nota_Export",
        body="# Cuerpo\n\nTexto canónico.\n",
        status="approved",
    )
    return document_id


def test_bridge_export_matches_canonical_note_document(temp_vault_manager):
    document_id = _prepare_approved_note(temp_vault_manager)
    backend = FunesConsoleBackend(temp_vault_manager.config.vault_path)
    backend.vault = temp_vault_manager
    bridge = FunesPyWebViewApi(backend)
    canonical = backend.get_notes_service().get_note(document_id).to_markdown()

    result = bridge.export_note(document_id, "markdown")

    assert "error" not in result
    assert result["source"] == "canonical"
    assert result["content"] == canonical
    metadata, body = parse_frontmatter(result["content"])
    assert metadata["status"] == "approved"
    assert body.startswith("# Cuerpo")


def test_export_service_and_bridge_agree_on_markdown(temp_vault_manager):
    vault_manager = temp_vault_manager
    vault_manager.create_theme(THEME)
    vault_manager.create_issue_in_theme(ISSUE)
    document_id, _ = write_note_under_theme(
        vault_manager,
        theme=THEME,
        issue=ISSUE,
        title="Nota_Acuerdo",
        body="# Acuerdo\n\nContenido idéntico.\n",
        status="approved",
    )
    resolver = AuthorizedPathResolver(
        vault_root=vault_manager.config.vault_path,
        output=vault_manager.output_dir,
        input=vault_manager.input_dir,
        dirty=vault_manager.dirty_dir,
        clean=vault_manager.clean_dir,
        quarantine=vault_manager.quarantine_dir,
    )
    store = JobStore(vault_manager.config.vault_path)
    notes = NotesApplicationService(
        vault=vault_manager,
        path_resolver=resolver,
        job_store=store,
        chroma_store=None,
    )
    export_service = ExportApplicationService(
        notes_service=notes,
        path_resolver=resolver,
    )
    try:
        service_payload = export_service.prepare_download(document_id, "markdown")
        backend = FunesConsoleBackend(vault_manager.config.vault_path)
        backend.vault = vault_manager
        bridge_payload = FunesPyWebViewApi(backend).export_note(document_id, "markdown")
        note = notes.get_note(document_id)
        assert service_payload.content == note.to_markdown()
        assert bridge_payload["content"] == service_payload.content
    finally:
        store.close()


def test_bridge_docx_export_round_trips_real_docx(temp_vault_manager):
    document_id = _prepare_approved_note(temp_vault_manager)
    backend = FunesConsoleBackend(temp_vault_manager.config.vault_path)
    backend.vault = temp_vault_manager
    result = FunesPyWebViewApi(backend).export_note(document_id, "docx")

    assert "error" not in result
    assert result["filename"].endswith(".docx")
    raw = base64.b64decode(result["content_base64"])
    assert raw.startswith(b"PK")

    document = Document(io.BytesIO(raw))
    note = backend.get_notes_service().get_note(document_id)
    assert "Nota_Export" in [paragraph.text for paragraph in document.paragraphs]
    assert any(
        paragraph.style.name == "Heading 1" and paragraph.text == "Cuerpo"
        for paragraph in document.paragraphs
    )
    assert any(
        "Texto canónico." in paragraph.text for paragraph in document.paragraphs
    )
    assert all(len(row.cells) == 2 for row in document.tables[0].rows)
    assert [row.cells[0].text for row in document.tables[0].rows] == sorted(
        note.frontmatter
    )
    metadata_rows = {
        row.cells[0].text: row.cells[1].text for row in document.tables[0].rows
    }
    assert metadata_rows["tags"] == json.dumps(
        note.frontmatter["tags"], ensure_ascii=False, sort_keys=True
    )
    assert note.to_markdown() not in [paragraph.text for paragraph in document.paragraphs]


def test_bridge_docx_keeps_code_and_unsupported_markdown_literal(temp_vault_manager):
    temp_vault_manager.create_theme(THEME)
    temp_vault_manager.create_issue_in_theme(ISSUE)
    body = (
        "Antes\n"
        "continúa\n\n"
        "```python\n"
        "print(\"hola\")\n"
        "```\n\n"
        "| A | B |\n"
        "| --- | --- |\n"
        "| 1 | 2 |\n\n"
        "> bloque literal\n\n"
        "#### Nivel cuatro literal\n"
    )
    document_id, _ = write_note_under_theme(
        temp_vault_manager,
        theme=THEME,
        issue=ISSUE,
        title="Nota_7B",
        body=body,
        status="approved",
    )

    backend = FunesConsoleBackend(temp_vault_manager.config.vault_path)
    backend.vault = temp_vault_manager
    result = FunesPyWebViewApi(backend).export_note(document_id, "docx")
    assert "error" not in result

    document = Document(io.BytesIO(base64.b64decode(result["content_base64"])))
    paragraphs = document.paragraphs
    assert "Antes\ncontinúa" in [paragraph.text for paragraph in paragraphs]

    code_paragraph = next(
        paragraph for paragraph in paragraphs if 'print("hola")' in paragraph.text
    )
    assert code_paragraph.style.name == "No Spacing"
    assert any(run.font.name == "Courier New" for run in code_paragraph.runs)
    assert "| A | B |\n| --- | --- |\n| 1 | 2 |" in [
        paragraph.text for paragraph in paragraphs
    ]
    assert "> bloque literal" in [paragraph.text for paragraph in paragraphs]
    assert "#### Nivel cuatro literal" in [
        paragraph.text for paragraph in paragraphs
    ]


def test_backend_approve_and_export_returns_review_result_without_destination_path(
    temp_vault_manager,
):
    temp_vault_manager.create_theme(THEME)
    temp_vault_manager.create_issue_in_theme(ISSUE)
    document_id, _ = write_note_under_theme(
        temp_vault_manager,
        theme=THEME,
        issue=ISSUE,
        title="Nota_9B",
        body="# Cuerpo 9B\n",
    )
    backend = FunesConsoleBackend(temp_vault_manager.config.vault_path)
    backend.vault = temp_vault_manager
    revision = backend.get_notes_service().get_note(document_id).revision

    result = backend.approve_and_export(
        document_id,
        revision,
        "markdown",
        metadata_patch={"title": "  Título actualizado  "},
    )

    assert result["approval_status"] == "approved"
    assert result["approved_revision"] == revision + 1
    assert result["export_status"] == "prepared"
    assert result["export_payload"]["source"] == "canonical"
    assert result["export_payload"]["content"].startswith("---\n")
    assert "destination_path" not in result
    assert "path" not in result["export_payload"]


def test_bridge_approve_and_export_normalizes_typed_metadata_before_backend(
    temp_vault_path,
):
    bridge = FunesPyWebViewApi(FunesConsoleBackend(temp_vault_path))
    calls = []
    bridge.backend.approve_and_export = (
        lambda document_id, expected_revision, export_format, metadata_patch=None: (
            calls.append(
                (document_id, expected_revision, export_format, metadata_patch)
            )
            or {"approval_status": "approved"}
        )
    )

    result = bridge.approve_and_export(
        "doc-id", 1, "markdown", metadata_patch={"title": "  Título  "}
    )

    assert result == {"approval_status": "approved"}
    assert calls == [("doc-id", 1, "markdown", {"title": "Título"})]


@pytest.mark.parametrize(
    "arguments,expected",
    [
        (
            ("folder/note.md", 1, "markdown", None),
            {"error": "path_not_authorized", "message": "Path is not authorized"},
        ),
        (
            ("doc-id", 0, "markdown", None),
            {
                "error": "invalid_payload",
                "message": "expected_revision must be a positive integer",
            },
        ),
        (
            ("doc-id", True, "markdown", None),
            {
                "error": "invalid_payload",
                "message": "expected_revision must be a positive integer",
            },
        ),
        (
            ("doc-id", 1, "rtf", None),
            {"error": "invalid_payload", "message": "format is not supported"},
        ),
        (
            ("doc-id", 1, "markdown", []),
            {"error": "invalid_payload", "message": "metadata_patch must be an object"},
        ),
    ],
)
def test_bridge_approve_and_export_rejects_invalid_boundary_inputs(
    arguments, expected, temp_vault_path
):
    bridge = FunesPyWebViewApi(FunesConsoleBackend(temp_vault_path))
    backend_calls = []
    bridge.backend.approve_and_export = lambda *args, **kwargs: backend_calls.append(
        (args, kwargs)
    )

    assert bridge.approve_and_export(*arguments) == expected
    assert backend_calls == []


def test_bridge_approve_and_export_maps_metadata_validation_to_stable_error(
    temp_vault_path,
):
    bridge = FunesPyWebViewApi(FunesConsoleBackend(temp_vault_path))
    result = bridge.approve_and_export(
        "doc-id", 1, "markdown", metadata_patch={"tags": ["bad:\nstatus: approved"]}
    )

    assert result["error"] == "invalid_metadata"
    assert "tags" in result["field_errors"]


def test_bridge_approve_and_export_maps_revision_conflict_to_stable_error(
    temp_vault_path,
):
    bridge = FunesPyWebViewApi(FunesConsoleBackend(temp_vault_path))

    def conflict(*args, **kwargs):
        raise NoteRevisionConflictError("doc-id")

    bridge.backend.approve_and_export = conflict

    assert bridge.approve_and_export("doc-id", 1, "markdown") == {
        "error": "note_revision_conflict",
        "message": "Note revision conflict: doc-id",
    }


def test_approval_export_contract_keeps_browser_projection_outside_backend_destination():
    source = CONSOLA_HTML.read_text(encoding="utf-8")

    assert "approve_and_export(currentSelectedDocumentId" in source
    assert "export_payload" in source
    assert "downloadTextBlob" in source
    assert "downloadBase64Blob" in source
    assert "openUserAssistedPdfPrintWindow" in source
    assert "completeUserAssistedPdfPrint" in source
    assert "destination_path" not in source[source.index("approveAndExportSelectedNote"):]
