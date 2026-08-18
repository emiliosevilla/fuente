"""Canonical note export service (Task 6.4)."""
from __future__ import annotations

import base64
import html
import io
import json
from pathlib import Path
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
import pytest

from fuente.application.export import (
    ExportApplicationService,
    ExportFileExistsError,
    UnsupportedExportFormatError,
)
from fuente.application.notes import NotesApplicationService
from fuente.domain.errors import (
    OutputApprovalRequiredError,
    PathAuthorizationError,
)
from fuente.domain.frontmatter import parse_frontmatter
from fuente.domain.paths import AuthorizedPathResolver
from fuente.infrastructure.sqlite_store import JobStore
from tests.conftest import approved_clean_origin, save_v3_summary_note


def _write_note(
    vault_manager,
    *,
    body: str,
    title: str = "Nota_Export",
    extra_metadata: dict | None = None,
) -> tuple[str, Path]:
    store = JobStore(vault_manager.config.vault_path)
    try:
        origin = approved_clean_origin(
            vault_manager,
            store,
            filename="origen-export.md",
        )
        return save_v3_summary_note(
            vault_manager,
            title=title,
            body=body,
            metadata_title=title.replace("_", " "),
            status="approved",
            tags=["export"],
            origins=[origin],
            extra_metadata=extra_metadata,
            store=store,
        )
    finally:
        store.close()


@pytest.fixture
def export_stack(temp_vault_manager):
    resolver = AuthorizedPathResolver(
        vault_root=temp_vault_manager.config.vault_path,
        output=temp_vault_manager.output_dir,
        input=temp_vault_manager.input_dir,
        dirty=temp_vault_manager.dirty_dir,
        clean=temp_vault_manager.clean_dir,
        quarantine=temp_vault_manager.quarantine_dir,
    )
    store = JobStore(temp_vault_manager.config.vault_path)
    notes = NotesApplicationService(
        vault=temp_vault_manager,
        path_resolver=resolver,
        job_store=store,
        chroma_store=None,
    )
    export_service = ExportApplicationService(
        notes_service=notes,
        path_resolver=resolver,
    )
    try:
        yield export_service, notes, temp_vault_manager, resolver
    finally:
        store.close()


def test_markdown_export_matches_canonical_document(export_stack):
    export_service, _, vault_manager, _ = export_stack
    body = "# Cuerpo\n\nTexto con <script>alert(1)</script> y [[wikilink]].\n"
    document_id, _ = _write_note(vault_manager, body=body)

    payload = export_service.prepare_download(document_id, "markdown")
    note = export_service.notes_service.get_note(document_id)

    assert payload.format == "markdown"
    assert payload.source == "canonical"
    assert payload.content == note.to_markdown()
    metadata, exported_body = parse_frontmatter(payload.content or "")
    assert metadata["title"] == "Nota Export"
    assert exported_body == body


def test_docx_export_is_real_docx_not_html_doc(export_stack):
    export_service, _, vault_manager, _ = export_stack
    document_id, _ = _write_note(vault_manager, body="# DOCX\n\nContenido.\n")

    payload = export_service.prepare_download(document_id, "word")
    raw = payload.content_bytes or b""

    assert payload.filename.endswith(".docx")
    assert raw.startswith(b"PK")
    assert b"[Content_Types].xml" in raw or b"word/" in raw


def test_docx_projection_is_byte_deterministic(export_stack):
    export_service, _, vault_manager, _ = export_stack
    document_id, _ = _write_note(vault_manager, body="# DOCX\n\nContenido.\n")
    with patch(
        "zipfile.time.localtime",
        return_value=(2026, 8, 18, 10, 0, 0, 1, 230, 0),
    ):
        first = export_service.prepare_download(document_id, "docx")
    with patch(
        "zipfile.time.localtime",
        return_value=(2026, 8, 18, 10, 0, 4, 1, 230, 0),
    ):
        second = export_service.prepare_download(document_id, "docx")
    assert first.content_bytes == second.content_bytes
    written_infos = []
    original_writestr = ZipFile.writestr

    def record_writestr(self, zinfo_or_arcname, data, compress_type=None):
        written_infos.append(zinfo_or_arcname)
        return original_writestr(self, zinfo_or_arcname, data, compress_type)

    with patch.object(ZipFile, "writestr", autospec=True, side_effect=record_writestr):
        canonical = ExportApplicationService._canonicalize_docx(
            first.content_bytes or b""
        )

    with ZipFile(io.BytesIO(first.content_bytes or b"")) as archive:
        entries = archive.infolist()
    with ZipFile(io.BytesIO(canonical)) as canonical_archive:
        canonical_entries = canonical_archive.infolist()
    assert entries
    assert len(written_infos) == len(entries)
    assert all(info.compress_type == ZIP_DEFLATED for info in written_infos)
    assert all(info.compress_level == 9 for info in written_infos)
    assert all(info.create_system == 3 for info in written_infos)
    assert all(info.external_attr == 0o600 << 16 for info in written_infos)
    assert all(info.compress_type == ZIP_DEFLATED for info in canonical_entries)
    assert all(info.create_system == 3 for info in canonical_entries)
    assert all(info.external_attr == 0o600 << 16 for info in canonical_entries)
    Document(io.BytesIO(first.content_bytes or b""))


def test_docx_projects_canonical_metadata_and_body_structures(export_stack):
    export_service, _, vault_manager, _ = export_stack
    body = (
        "# Resumen Ejecutivo\n\n"
        "Primer párrafo\ncontinúa en la misma sección.\n\n"
        "## Desarrollo\n\n"
        "### Detalle\n\n"
        "- elemento\n"
        "1. paso\n"
    )
    document_id, _ = _write_note(
        vault_manager, body=body, title="Nota Estructurada"
    )
    note = export_service.notes_service.get_note(document_id)

    payload = export_service.prepare_download(document_id, "docx")
    document = Document(io.BytesIO(payload.content_bytes or b""))
    texts = [paragraph.text for paragraph in document.paragraphs]

    assert note.title in texts
    assert any(note.relative_path in text for text in texts)
    assert "Resumen Ejecutivo" in texts
    assert "Primer párrafo\ncontinúa en la misma sección." in texts
    assert any(
        paragraph.style.name == "Heading 2" and paragraph.text == "Desarrollo"
        for paragraph in document.paragraphs
    )
    assert any(
        paragraph.style.name == "Heading 3" and paragraph.text == "Detalle"
        for paragraph in document.paragraphs
    )
    assert any(
        paragraph.style.name == "List Bullet" and paragraph.text == "elemento"
        for paragraph in document.paragraphs
    )
    assert any(
        paragraph.style.name == "List Number" and paragraph.text == "paso"
        for paragraph in document.paragraphs
    )
    assert all(len(row.cells) == 2 for row in document.tables[0].rows)
    assert [row.cells[0].text for row in document.tables[0].rows] == sorted(
        note.frontmatter
    )
    metadata_rows = {
        row.cells[0].text: row.cells[1].text for row in document.tables[0].rows
    }
    assert metadata_rows["tags"] == json.dumps(
        note.frontmatter["tags"], ensure_ascii=False, sort_keys=True
    )
    assert metadata_rows["history"] == json.dumps(
        note.frontmatter["history"], ensure_ascii=False, sort_keys=True
    )
    assert note.to_markdown() not in texts


def test_docx_serializes_dict_metadata_as_sorted_json(export_stack):
    export_service, _, vault_manager, _ = export_stack
    extra_metadata = {"custom_metadata": {"zeta": 2, "alpha": 1}}
    document_id, _ = _write_note(
        vault_manager,
        body="# Metadata\n",
        extra_metadata=extra_metadata,
    )

    note = export_service.notes_service.get_note(document_id)
    payload = export_service.prepare_download(document_id, "docx")
    document = Document(io.BytesIO(payload.content_bytes or b""))
    metadata_rows = {
        row.cells[0].text: row.cells[1].text for row in document.tables[0].rows
    }

    assert all(len(row.cells) == 2 for row in document.tables[0].rows)
    assert metadata_rows["custom_metadata"] == '{"alpha": 1, "zeta": 2}'
    assert metadata_rows["custom_metadata"] == json.dumps(
        note.frontmatter["custom_metadata"], ensure_ascii=False, sort_keys=True
    )


def test_docx_preserves_paragraphs_code_and_unsupported_markdown(export_stack):
    export_service, _, vault_manager, _ = export_stack
    table = "| Col A | Col B |\n| --- | --- |\n| uno | dos |"
    body = (
        "Primera línea\n"
        "segunda línea\n\n"
        "Párrafo separado\n\n"
        "```python\n"
        "def saludo():\n"
        "    return \"hola\"\n"
        "\n"
        "```\n\n"
        f"{table}\n\n"
        "> cita literal\n\n"
        "#### Heading literal\n"
    )
    document_id, _ = _write_note(vault_manager, body=body)

    payload = export_service.prepare_download(document_id, "docx")
    document = Document(io.BytesIO(payload.content_bytes or b""))
    paragraphs = document.paragraphs

    assert "Primera línea\nsegunda línea" in [paragraph.text for paragraph in paragraphs]
    assert "Párrafo separado" in [paragraph.text for paragraph in paragraphs]

    code_paragraph = next(
        paragraph for paragraph in paragraphs if "def saludo():" in paragraph.text
    )
    assert code_paragraph.style.name == "No Spacing"
    assert code_paragraph.text == 'def saludo():\n    return "hola"\n'
    assert any(
        run.font.name == "Courier New" for run in code_paragraph.runs
    )

    texts = [paragraph.text for paragraph in paragraphs]
    assert table in texts
    assert "> cita literal" in texts
    assert "#### Heading literal" in texts
    assert all(
        fragment in "\n".join(texts)
        for fragment in ("| Col A | Col B |", "| --- | --- |", "| uno | dos |")
    )


def test_pdf_export_is_user_assisted_with_escaped_html(export_stack):
    export_service, _, vault_manager, _ = export_stack
    body = "# Título\n\n<script>x</script>\n"
    document_id, note_path = _write_note(vault_manager, body=body)
    relative = note_path.resolve().relative_to(
        vault_manager.config.vault_path.resolve()
    ).as_posix()
    note = export_service.notes_service.get_note(document_id)

    payload = export_service.prepare_download(document_id, "pdf")

    assert payload.mode == "user_assisted_print"
    assert "impresión asistida" in (payload.label or "").lower()
    assert "contenido canónico" in (payload.label or "").lower()
    assert "<script>" not in payload.print_html
    assert "&lt;script&gt;" in payload.print_html
    assert relative in payload.print_html
    assert "pdf-frontmatter" in payload.print_html
    assert "status: approved" in payload.print_html
    assert "- export" in payload.print_html
    assert html.escape(note.title) in payload.print_html
    assert "<h1>Título</h1>" in payload.print_html


def test_write_export_rejects_unauthorized_destination(export_stack):
    export_service, _, vault_manager, _ = export_stack
    document_id, _ = _write_note(vault_manager, body="# Nota\n")

    with pytest.raises(PathAuthorizationError):
        export_service.write_export(
            document_id,
            "markdown",
            "../outside/nota.md",
        )


def test_write_export_blocks_overwrite_without_confirmation(export_stack):
    export_service, _, vault_manager, _ = export_stack
    document_id, _ = _write_note(vault_manager, body="# Nota\n")
    existing = vault_manager.output_dir / "existing_export.md"
    existing.write_text("occupied", encoding="utf-8")
    existing_relative = existing.resolve().relative_to(
        vault_manager.config.vault_path.resolve()
    ).as_posix()

    with pytest.raises(ExportFileExistsError):
        export_service.write_export(
            document_id,
            "markdown",
            existing_relative,
            confirm_overwrite=False,
        )

    result = export_service.write_export(
        document_id,
        "markdown",
        existing_relative,
        confirm_overwrite=True,
    )

    assert result["status"] == "exported"
    assert existing.read_text(encoding="utf-8") != "occupied"


def test_unsupported_format_raises_stable_error(export_stack):
    export_service, _, vault_manager, _ = export_stack
    document_id, _ = _write_note(vault_manager, body="# Nota\n")

    with pytest.raises(UnsupportedExportFormatError) as raised:
        export_service.prepare_download(document_id, "rtf")

    assert raised.value.code == "unsupported_export_format"


def test_console_export_note_returns_canonical_payload(temp_vault_manager):
    from fuente.control_console import FuenteConsoleBackend

    document_id, _ = _write_note(temp_vault_manager, body="# Backend\n")
    backend = FuenteConsoleBackend(temp_vault_manager.config.vault_path)

    result = backend.export_note(document_id, "markdown")

    assert result["source"] == "canonical"
    assert "content" in result
    metadata, body = parse_frontmatter(result["content"])
    assert metadata["status"] == "approved"
    assert body.startswith("# Backend")


def test_pending_output_cannot_be_exported_before_editorial_approval(export_stack):
    export_service, notes, vault_manager, _resolver = export_stack
    origin = approved_clean_origin(
        vault_manager,
        notes.job_store,
        filename="origen-export-pending.md",
    )
    document_id, _ = save_v3_summary_note(
        vault_manager,
        title="Salida_Pendiente",
        body="# Pendiente\n",
        status="pending_review",
        origins=[origin],
        store=notes.job_store,
    )

    with pytest.raises(OutputApprovalRequiredError):
        export_service.prepare_download(document_id, "markdown")


def test_docx_payload_round_trips_via_base64_dict(export_stack):
    export_service, _, vault_manager, _ = export_stack
    document_id, _ = _write_note(vault_manager, body="# Base64\n")

    payload = export_service.prepare_download(document_id, "docx").as_dict()

    assert payload["content_base64"]
    decoded = base64.b64decode(payload["content_base64"])
    assert decoded.startswith(b"PK")
